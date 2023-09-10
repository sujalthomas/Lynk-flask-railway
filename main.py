from flask import Flask, render_template
from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_mail import Mail as FlaskMail, Message 
from flask_session import Session
from flask_bcrypt import Bcrypt
import os , random , secrets
from dotenv import load_dotenv
import redis
import PyPDF2
from docx import Document 
from flask_cors import CORS
import logging 
from itsdangerous import URLSafeTimedSerializer
from flask_limiter import Limiter
import logging
from flask_limiter.util import get_remote_address
import openai
from datetime import datetime, timedelta
from werkzeug.security import generate_password_hash as encrypt_password
from flask_mail import Message
from http import HTTPStatus
import re , io
from werkzeug.security import generate_password_hash, check_password_hash
from wtforms import Form, StringField, PasswordField, validators


# Load environment variables
load_dotenv()

app = Flask(__name__)


bcrypt = Bcrypt(app)


# Secuirty configurations
SECRET_KEY = os.getenv("SECRET_KEY", default=secrets.token_urlsafe(16))
app.config["SECRET_KEY"] = SECRET_KEY
serializer = URLSafeTimedSerializer(app.config["SECRET_KEY"])

app.config["SECURITY_PASSWORD_SALT"] = os.getenv(
    "SECURITY_PASSWORD_SALT", default="your_random_salt"
)
app.config["SECURITY_REGISTERABLE"] = True
app.config["SECURITY_RECOVERABLE"] = True
app.config["SECURITY_PASSWORD_SALT"] = os.getenv(
    "SECURITY_PASSWORD_SALT", default="your_random_salt"
)

# Database configurations
DB_USERNAME = os.getenv("DB_USERNAME")
DB_PASSWORD = os.getenv("DB_PASSWORD")
DB_HOST = os.getenv("DB_HOST")
DB_PORT = os.getenv("DB_PORT")
DB_NAME = os.getenv("DB_NAME")


app.config[
    "SQLALCHEMY_DATABASE_URI"
] = f"postgresql://postgres:69qK1AyQZVbWpUxpZ1EZ@containers-us-west-105.railway.app:5544/railway"


# Database configuration
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)



CORS(
    app,
    resources={r"/*": {"origins": ["https://www.linkedin.com"]}},
    supports_credentials=True,
)

# Mail Configurations
app.config["MAIL_SERVER"] = "smtp.sendgrid.net"
app.config["MAIL_PORT"] = 587  # 465 for TLS
app.config["MAIL_USERNAME"] = "apikey"
app.config["MAIL_PASSWORD"] = "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"
app.config["MAIL_USE_TLS"] = True
app.config["MAIL_USE_SSL"] = False
mail = FlaskMail(app)


# Session configurations
sess = Session()
sess.init_app(app)
mail = FlaskMail(app)


# Session configurations 
app.config["REDIS_URL"] = "redis://default:S5rZfd5YEDm88llfugC5@containers-us-west-154.railway.app:7407"

REDIS_URL = "redis://default:S5rZfd5YEDm88llfugC5@containers-us-west-154.railway.app:7407"


limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["5 per minute"],
    storage_uri=REDIS_URL
)


# Generate DB tables
with app.app_context():
    db.create_all()




# Current folder path
UPLOAD_FOLDER = os.path.dirname(os.path.abspath(__file__))

# Configuration options
logging.basicConfig(level=logging.INFO)



# MODELS ###############
# Define models #######

#cvs generated
class CoverLetter(db.Model):
    cover_letter_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.user_id"), nullable=False)
    company_name = db.Column(db.String(100))
    job_listing = db.Column(db.Text)
    recruiter = db.Column(db.String(100))
    date = db.Column(db.DateTime, default=db.func.CURRENT_TIMESTAMP)
    file_path = db.Column(db.String(255))
    user = db.relationship("User", backref="cover_letters")

# resumes uploaded
class Resume(db.Model):
    resume_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.user_id"), nullable=False)
    content = db.Column(db.Text)
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    user = db.relationship("User", backref="resumes")

# usage statistics
class UsageStatistic(db.Model):
    stat_id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.user_id"))
    action = db.Column(db.String(50))
    date = db.Column(db.DateTime, default=db.func.CURRENT_TIMESTAMP)
    user = db.relationship("User", backref="usage_statistics")

# Define roles
class Role(db.Model):
    id = db.Column(db.Integer(), primary_key=True)
    name = db.Column(db.String(80), unique=True)
    description = db.Column(db.String(255))


# user roles
roles_users = db.Table(
    "roles_users",
    db.Column("user_id", db.Integer(), db.ForeignKey("user.user_id")),
    db.Column("role_id", db.Integer(), db.ForeignKey("role.id")),
)

class User(db.Model):
    user_id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(255))
    password = db.Column(db.String(255), nullable=False)
    email = db.Column(db.String(150), unique=True, nullable=False)
    fs_uniquifier = db.Column(db.String(255), unique=True)
    roles = db.relationship(
        "Role", secondary=roles_users, backref=db.backref("users", lazy="dynamic")
    )
    
    # New columns for reset code and its expiration
    password_reset_code = db.Column(db.String(6))
    password_reset_code_expiration = db.Column(db.DateTime)
    
    # New columns for email verification
    # New columns for authentication and authorization
    is_active = db.Column(db.Boolean, default=False)
    verification_code = db.Column(db.String(6))
    verification_code_expiration = db.Column(db.DateTime)

    # New methods and properties for authentication and authorization
    @property
    def active(self):
        return self.is_active

    def get_id(self):
        return self.user_id

    def has_role(self, role):
        return role in [role.name for role in self.roles]

# Define blacklisted tokens
class BlacklistedToken(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    token = db.Column(db.String(500), unique=True, nullable=False)
    blacklisted_on = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

    @staticmethod
    def check_blacklist(auth_token):
        res = BlacklistedToken.query.filter_by(token=str(auth_token)).first()
        if res:
            return True
        else:
            return False




# Define user data store
class UserDatastore:
    def __init__(self, db, user_model, role_model):
        self.db = db
        self.user_model = user_model
        self.role_model = role_model

    def find_user(self, **kwargs):
        return self.user_model.query.filter_by(**kwargs).first()

    def find_role(self, role):
        return self.role_model.query.filter_by(name=role).first()

    def add_role_to_user(self, user, role):
        role = self.find_role(role)
        if role:
            user.roles.append(role)
            self.db.session.commit()

    def remove_role_from_user(self, user, role):
        role = self.find_role(role)
        if role:
            user.roles.remove(role)
            self.db.session.commit()

    def get_user(self, user_id):
        return self.user_model.query.get(user_id)

    def get_role(self, role_id):
        return self.role_model.query.get(role_id)



# Define hashed_password
password = "supersecretpassword"
hashed_password = encrypt_password(password)

# Setup Flask-Security
user_datastore = UserDatastore(db, User, Role)


# Custom registration form
class RegistrationForm(Form):
    username = StringField('Username', [validators.DataRequired()])
    password = PasswordField('Password', [validators.DataRequired()])
    # ...
    # any other fields and validators that you need
    # ...


# Utils
# Define utils
def is_api_key_valid(api_key):
    openai.api_key = api_key
    try:
        response = openai.Completion.create(
            engine="davinci", prompt="This is a test.", max_tokens=5
        )
    except:
        return False
    else:
        return True

def convert_to_txt(file, file_type):
    if file_type == "docx":
        doc = Document(file)
        data = "\n".join([p.text for p in doc.paragraphs])
    elif file_type == "pdf":
        reader = PyPDF2.PdfReader(file)
        data = "\n".join(
            [reader.pages[i].extract_text() for i in range(len(reader.pages))]
        )
    else:
        raise ValueError("Unsupported file type")
    return data.replace('\0', '')




# Routes ###############
# Define routes #######
# api key verification


@app.route('/')
def index():
    return render_template('index.html', title='Home' )

@app.route('/google3a556bba71bd41e1.html')
def google():
    return render_template('google3a556bba71bd41e1.html', title='Home' )
    

# Routes ###############
# Define routes #######
# api key verification
@app.route("/apiverify", methods=["POST"])
def apiverify():
    data = request.get_json()
    api_key = data.get("apiKey")

    try:
        if is_api_key_valid(api_key):
            token = serializer.dumps({"user": "YOUR_USER_IDENTIFIER"})
            return jsonify(success=True, token=token)
        else:
            raise Exception("Invalid API key")
    except Exception as e:
        logging.error(str(e))
        return jsonify(success=False, message="Invalid API key"), 401



# user registration
@app.route("/request-reset-password", methods=["POST"])
@limiter.limit(
    "2 per minute"
)  # For instance, limit to 2 requests per minute for this route
def request_reset_password():
    data = request.get_json()
    logging.info(f"Hit /request-reset-password with data: {data}")
    email = data.get("email")
    user = User.query.filter_by(email=email).first()
    if not user:
        logging.warning(f"Password reset attempt for non-existent email: {email}")
        return jsonify(success=False, message="User not found"), 404

    code = str(random.randint(100000, 999999))
    user.password_reset_code = code
    user.password_reset_code_expiration = datetime.utcnow() + timedelta(minutes=30)
    db.session.commit()

    msg = Message(
        "Password Reset Code", sender="lynktools@gmail.com", recipients=[email]
    )

    msg.body = f"Your password reset code is: {code}"
    try:
        mail.send(msg)
    except Exception as e:
        logging.error(f"Error sending email to {email}: {str(e)}")
        return jsonify(success=False, message="Error sending email."), 500

    logging.info(f"Sent password reset code to: {email}")
    return jsonify(success=True, message="Password reset code has been sent."), 200



@app.route("/reset-password", methods=["POST"])
@limiter.limit(
    "3 per minute"
)  # For instance, limit to 3 requests per minute for this route
def reset_password_with_code():
    data = request.get_json()
    logging.info(f"Hit /reset-password with data: {data}")
    email = data.get("email")
    code = data.get("code")
    new_password = data.get("newPassword")

    user = User.query.filter_by(email=email).first()

    if not user:
        logging.warning(f"Password reset attempt for non-existent email: {email}")
        return jsonify(success=False, message="User not found"), 404

    if (
        user.password_reset_code != code
        or datetime.utcnow() > user.password_reset_code_expiration
    ):
        logging.warning(f"Invalid or expired code used for email: {email}")
        return jsonify(success=False, message="Invalid or expired code"), 401

    hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
    user.password = hashed_password

    user.password_reset_code = None
    user.password_reset_code_expiration = None
    db.session.commit()

    logging.info(f"Password reset successful for: {email}")
    return jsonify(success=True, message="Password reset successful"), 200



@app.route("/register", methods=["POST"])
def register():
    data = request.get_json()
    email = data.get("email")
    password = data.get("password")
    name = data.get("name")

    hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

    user = User.query.filter_by(email=email).first()
    if user:
        return jsonify(success=False, message="User already exists"), 400

    user_role = Role.query.filter_by(name="user").first()
    if not user_role:
        user_role = Role(name="user", description="Regular user")
        db.session.add(user_role)
        db.session.commit()

    new_user = User(email=email, password=hashed_password, is_active=False)
    new_user.roles.append(user_role)
    db.session.add(new_user)
    db.session.commit()

    # Generate verification code
    code = str(random.randint(100000, 999999))
    new_user.verification_code = code
    new_user.verification_code_expiration = datetime.utcnow() + timedelta(minutes=30)
    db.session.commit()

    # Send verification code to user's email
    msg = Message("Verification Code", sender="lynktools@gmail.com", recipients=[email])
    msg.body = f"Your verification code is: {code}"
    try:
        mail.send(msg)
    except Exception as e:
        logging.error(f"Error sending email to {email}: {str(e)}")
        return jsonify(success=False, message="Error sending email."), 500

    return jsonify(success=True, message="Verification code has been sent to your email."), 200

@app.route("/verify", methods=["POST"])
def verify():
    data = request.get_json()
    email = data.get("email")
    code = data.get("code")

    user = User.query.filter_by(email=email).first()
    if not user:
        return jsonify(success=False, message="User not found"), 404

    if user.verification_code != code or datetime.utcnow() > user.verification_code_expiration:
        return jsonify(success=False, message="Invalid or expired code"), 401

    user.is_active = True
    user.verification_code = None
    user.verification_code_expiration = None
    db.session.commit()

    return jsonify(success=True, message="Registration successful"), 200

@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    email = data.get("email")
    password = data.get("password")

    user = User.query.filter_by(email=email).first()
    
    if not user:
        return jsonify(success=False, message="No account found with this email address."), 404

    if not bcrypt.check_password_hash(user.password, password):
        return jsonify(success=False, message="Incorrect password."), 401

    if not user.is_active:
        return jsonify(success=False, message="Your account is not active. Please verify your email or contact support."), 403

    token = serializer.dumps({"user": user.user_id}, salt="password-reset")
    return jsonify(success=True, token=token, message="Login successful."), 200


# @app.route("/logout", methods=["POST"])
# def logout():
#     token = request.headers.get("Authorization")
#     try:
#         data = serializer.loads(token, salt="password-reset", max_age=9600)
#     except Exception as e:
#         print(f"Token deserialization error: {e}")
#         return jsonify(success=False, message="Invalid or expired token"), 401
    
#     data = request.get_json()
#     # Ensure the token belongs to a user
#     user_id = data.get("user_id")


#     decoded_data = serializer.loads(user_id, salt="password-reset", max_age=9600)

#     user_id_decode = decoded_data["user"]
#     if not user_id_decode:
#         return jsonify(success=False, message="User not found."), 404


#     # If the token is not already blacklisted, add it to the blacklist
#     if not BlacklistedToken.check_blacklist(token):
#         blacklist_token = BlacklistedToken(token=token)
#         try:
#             db.session.add(blacklist_token)
#             db.session.commit()
#         except Exception as e:
#             return jsonify(success=False, message=f"Failed to logout. {e}"), 500

#     return jsonify(success=True, message="Logout successful."), 200

@app.route("/logout", methods=["POST"])
def logout():
    # In the future, potential to invalidate tokens
    return jsonify(success=True, message="Logout successful."), 200




user_counters = {}


# cover letter generation
@app.route("/cover-letter", methods=["POST"])
def listen():
    token = request.headers.get("Authorization")
    try:
        data = serializer.loads(token, salt="password-reset", max_age=9600)
    except Exception as e:
        print(f"Token deserialization error: {e}")
        return jsonify(success=False, message="Invalid or expired token"), 401

    data = request.get_json()
    api_key = data.get("apiKey")
    openai.api_key = api_key

    # request body data from the frontend and assign it as user_id

    company_name = data.get("Company-name", "")
    job_listing = data.get("Job-Listing", "")
    recruiter = data.get("Recruiter", "")
    date = data.get("Date", "")
    user_id = data.get("user_id")

    decoded_data = serializer.loads(user_id, salt="password-reset", max_age=9600)

    user_id_decode = decoded_data["user"]
    print("User ID:", user_id_decode)

    # Construct the path to the user's folder and the resume inside
    user_folder_path = os.path.join(UPLOAD_FOLDER, str(user_id_decode))
    resume_path = os.path.join(user_folder_path, "current_resume.txt")

    try:
        with open(resume_path, "r") as file:
            resume = file.read()
    except FileNotFoundError:
        return jsonify(success=False, message="Resume file not found."), 404

    try:
        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant that will craft a tailored cover letter using a resume and a job listing. Your goal is to create a compelling cover letter that showcases the candidate's skills and experiences, aligning them with the job's requirements.",
                },
                {
                    "role": "user",
                    "content": f"Using this resume, {resume}, and this job listing, {job_listing}, craft a cover letter that doesn't include addresses but highlights the candidate's fit for the role. Ensure it includes the candidate's name, email, phone number, and LinkedIn profile. Also, only include the company name {company_name}, followed by recruiter's name {recruiter} and today's date {date}. No place holder text is allowed, if recruiters name is not found use 'Dear Hiring Manager.' ",
                },
            ],
            temperature=1.3,
            top_p=0.9,
            max_tokens=700,
            frequency_penalty=0.5,
            presence_penalty=0.5,
        )
    except openai.error.OpenAIError as e:
        print("OpenAI API Error:", e)
        return jsonify(success=False, message="OpenAI API Error"), 500

    cover_letter_content = completion.choices[0].message.content
    base_filename = f"{company_name}_cv.docx"
    filename = base_filename
    count = 1

    # Check if "Generated_CVs" folder exists, if not, create it
    folder_name = "Generated_CVs"
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)

    # Ensure unique filename within the "Generated_CVs" folder
    while os.path.isfile(os.path.join(folder_name, filename)):
        filename = f"{company_name}_cv({count}).docx"
        count += 1

    # Create the DOCX in memory
    doc = Document()
    doc.add_paragraph(cover_letter_content)
    mem_stream = io.BytesIO()
    doc.save(mem_stream)

    # Reset stream position
    mem_stream.seek(0)

    # Increment user's counter
    user_id_decode = str(user_id_decode)  # Ensure the key is a string
    user_counters[user_id_decode] = user_counters.get(user_id_decode, 0) + 1
    print(
        f"User {user_id_decode} has generated {user_counters[user_id_decode]} cover letters."
    )

    # Send the in-memory DOCX as a file
    return send_file(
        mem_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# resume generation
@app.route("/generate-resume", methods=["POST"])
def generate_resume():
    token = request.headers.get("Authorization")
    try:
        data = serializer.loads(token, salt="password-reset", max_age=9600)
    except Exception as e:
        print(f"Token deserialization error: {e}")
        return jsonify(success=False, message="Invalid or expired token"), 401

    data = request.get_json()
    api_key = data.get("apiKey")
    openai.api_key = api_key

    job_description = data.get("Job-Description", "")
    user_id = data.get("user_id")

    print(user_id)
    print(job_description)

    decoded_data = serializer.loads(user_id, salt="password-reset", max_age=9600)

    user_id_decode = decoded_data["user"]
    print("User ID:", user_id_decode)

    # Construct the path to the user's folder and the resume inside
    user_folder_path = os.path.join(UPLOAD_FOLDER, str(user_id_decode))
    resume_path = os.path.join(user_folder_path, "current_resume.txt")

    try:
        with open(resume_path, "r") as file:
            resume = file.read()
    except FileNotFoundError:
        return jsonify(success=False, message="Resume file not found."), 404

    print(resume)

    try:
        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant that will reword a resume to better align it with a job description. Your goal is to highlight the candidate's skills and experiences, making them match the job's requirements as closely as possible.",
                },
                {
                    "role": "user",
                    "content": f"Given this original resume: {resume}, and this job description: {job_description}, please reword the resume to better fit the job requirements.",
                },
            ],
            temperature=1.3,
            top_p=0.9,
            max_tokens=1000,
            frequency_penalty=0.5,
            presence_penalty=0.5,
        )
    except openai.error.OpenAIError as e:
        print("OpenAI API Error:", e)
        return jsonify(success=False, message="OpenAI API Error"), 500

    reworded_resume_content = completion.choices[0].message.content
    base_filename = "reworded_resume.docx"
    filename = base_filename
    count = 1

    # Check if "Generated_Resumes" folder exists, if not, create it
    folder_name = "Generated_Resumes"
    if not os.path.exists(folder_name):
        os.mkdir(folder_name)

    # Ensure unique filename within the "Generated_Resumes" folder
    while os.path.isfile(os.path.join(folder_name, filename)):
        filename = f"reworded_resume({count}).docx"
        count += 1

    # Create the DOCX in memory
    doc = Document()
    doc.add_paragraph(reworded_resume_content)
    mem_stream = io.BytesIO()
    doc.save(mem_stream)

    # Reset stream position
    mem_stream.seek(0)

    # Increment user's counter
    user_id_decode = str(user_id_decode)  # Ensure the key is a string
    user_counters[user_id_decode] = user_counters.get(user_id_decode, 0) + 1
    print(
        f"User {user_id_decode} has generated {user_counters[user_id_decode]} cover letters."
    )

    # Send the in-memory DOCX as a file
    return send_file(
        mem_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )




@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    # Retrieve the token from formData
    token = request.form.get("user_id")
    print(token)

    # Check if the token is provided
    if not token:
        return jsonify(success=False, message="Token not provided"), 421

    # Deserialize the token to extract user_id
    try:
        data = serializer.loads(token, salt="password-reset", max_age=9600)
        user_id = data["user"]
        print(user_id)
    except Exception as e:
        print(f"Deserialization error: {e}")
        return jsonify(success=False, message="Invalid or expired token"), 420

    # Check if file is present in the request
    if "resume" not in request.files:
        return jsonify(success=False, message="No file part"), 400

    file = request.files["resume"]

    # Check if a filename is provided
    if file.filename == "":
        return jsonify(success=False, message="No selected file"), 400

    # Check for allowed file extensions
    ALLOWED_EXTENSIONS = ["pdf", "docx"]
    if file.filename.rsplit(".", 1)[1].lower() not in ALLOWED_EXTENSIONS:
        return jsonify(success=False, message="File type not allowed"), 400

    # Ensure file is not too large
    if len(file.read()) > 5 * 1024 * 1024:
        return jsonify(success=False, message="File size exceeds the limit"), 400

    # Reset file pointer after reading
    file.seek(0)
    file_type = os.path.splitext(file.filename)[1][1:]

    # Convert file to txt
    try:
        txt_content = convert_to_txt(file, file_type)
    except ValueError:
        return jsonify(success=False, message="Unsupported file type"), 400
    except Exception as e:
        return jsonify(success=False, message=str(e)), 500

    # Store in database
    resume = Resume(user_id=user_id, content=txt_content)
    db.session.add(resume)
    db.session.commit()

    # Store the text content in the user's folder
    user_folder = os.path.join(UPLOAD_FOLDER, str(user_id))
    if not os.path.exists(user_folder):
        os.mkdir(user_folder)

    filename = os.path.join(user_folder, "current_resume.txt")
    with open(filename, "w", encoding="utf-8") as txt_file:
        txt_file.write(txt_content)

    return (
        jsonify(success=True, message="File uploaded and converted successfully"),
        200,
    )



if __name__ == '__main__':
    app.run(debug=True, port=os.getenv("PORT", default=5000))
